from io import BytesIO

from docx import Document


def docx_bytes(
    text: str = "This is a statement of work for the engagement.",
    table_rows: list[list[str]] | None = None,
) -> bytes:
    document = Document()
    if text:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row_values in enumerate(table_rows):
            for col_index, value in enumerate(row_values):
                table.rows[row_index].cells[col_index].text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_bytes(text: str = "This is a statement of work for the engagement.") -> bytes:
    if not text.strip():
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buffer = BytesIO()
        writer.write(buffer)
        return buffer.getvalue()
    return _pdf_with_text(text)


def _pdf_with_text(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", "replace")
    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n",
        (
            f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode()
            + stream
            + b"\nendstream\nendobj\n"
        ),
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]
    header = b"%PDF-1.4\n"
    offsets = [0]
    body = bytearray()
    cursor = len(header)
    for obj in objects:
        offsets.append(cursor)
        body.extend(obj)
        cursor += len(obj)
    xref_start = cursor
    xref = bytearray(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        xref.extend(f"{offset:010d} 00000 n \n".encode())
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n"
    ).encode()
    return bytes(header + body + xref + trailer)


def mpp_stub_bytes(payload: bytes = b"stub-project") -> bytes:
    """OLE compound header so FileValidator accepts the upload as an MPP type."""
    return b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + payload


def pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)
