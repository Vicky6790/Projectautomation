from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.errors import AppError

_MIN_TEXT_CHARS = 20


def extract_sow_text(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = _pdf_text(content)
    elif ext == ".docx":
        text = _docx_text(content)
    else:
        raise AppError(
            400,
            "UNSUPPORTED_FILE_TYPE",
            "SOW text extraction supports PDF and Word only",
        )
    cleaned = " ".join(text.split())
    if len(cleaned) < _MIN_TEXT_CHARS:
        raise AppError(
            400,
            "NO_EXTRACTABLE_TEXT",
            "No extractable text was found. Provide a text-based PDF or Word document, "
            "not a scanned or image-only file",
        )
    return cleaned


def _pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 - invalid PDF is a user error
        raise AppError(400, "UNSUPPORTED_FILE_TYPE", "The PDF could not be read") from exc


def _docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        for section in document.sections:
            for paragraph in (*section.header.paragraphs, *section.footer.paragraphs):
                if paragraph.text.strip():
                    parts.append(paragraph.text)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - invalid Word file is a user error
        raise AppError(400, "UNSUPPORTED_FILE_TYPE", "The Word document could not be read") from exc
